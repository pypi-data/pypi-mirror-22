from time import gmtime, strftime
from docx.shared import Inches
from docx.shared import Pt
import matplotlib.pyplot as plt
import seaborn as sns
import seaborn.linearmodels as snsl
import tushare as ts
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from timeit import default_timer as timer
from datetime import datetime, timedelta
import matplotlib.dates as mdates
from matplotlib.pyplot import subplots, draw
from matplotlib.finance import _candlestick
import pandas as pd
from pandas.tseries.offsets import BDay

def shiftFromToday(day, isBack): #int
    today = pd.datetime.today()
    if isBack:
        theday = today - BDay(day)
        print(theday)
        theday = '{}'.format(theday)
        theday = theday.split(' ')[0]
        return theday
    else:
        theday = today + BDay(day)
        print(theday)
        theday = '{}'.format(theday)
        theday = theday.split(' ')[0]
        return theday
def today():
    return '{}'.format(strftime("%Y-%m-%d", gmtime()))  # "%Y-%m-%d %H:%M:%S"
def yesterday():
    return '{}'.format(datetime.strftime(datetime.now() - timedelta(1), '%Y-%m-%d'))


def templeateForming(): #with layout
    # TEMPLATE FORMING
    template = "tushareUse/Files/input.docx"
    document = ts.MailMerge(template)
    document.merge(
        today=today(),
        p1='Springfield',
        p2='800-555-5555')
    document.write('tushareUse/Files/output.docx')

#Document editing
wordDocument = Document('tushareUse/Files/output.docx')

def newDocument():
    wordDocument._body.clear_content()

def openDocument():
    p = wordDocument.add_paragraph()
    setParagraphStyle(p)
    r = p.add_run()
    return r

def saveDocument():
    wordDocument.save('tushareUse/Files/output.docx')

def addTableToDocument(rows, cols):
    # styles = [s for s in wordDocument.styles if s.type == WD_STYLE_TYPE.TABLE]
    # for style in styles:
    #     print(style.name)
    table = wordDocument.add_table(rows = rows, cols = cols)
    table.style = 'Light Shading Accent 3'
    return table

def setParagraphStyle(p):
    style = wordDocument.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(15)
    p.style = wordDocument.styles['Normal']

def setPageMargin(inputfile): ## not tested
    # Open the document
    document = Document(outputfile)
    # changing the page margins
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(0)
        section.right_margin = Cm(0)

    document.save(args.outputFile)

def drawCandlestick(df):
    #weekFormatter = DateFormatter('%b %d')     # 如：Jan 12
    fig, ax = plt.subplots()
    fig.subplots_adjust(bottom=0.2)
    df['t'] = df['date']
    _candlestick(ax, df, width=0.6, colorup='r', colordown='g')

    ax.xaxis_date()
    ax.autoscale_view()
    plt.setp(plt.gca().get_xticklabels(), rotation=45, horizontalalignment='right')


    ax.grid(True)
    plt.title('000777', fontproperties=zhfont)
    plt.show()

def getCodeArray():
    allStocks = ts.get_today_all()
    codeArray = []
    for i in range(0, allStocks.shape[0]):  # pandas dataframe row count, column count is stock.shape[1]
        codeArray.append(allStocks.iloc[i]['code'])
    return codeArray